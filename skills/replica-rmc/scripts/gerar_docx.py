#!/usr/bin/env python3
"""gerar_docx.py — monta .docx final da replica RMC/RCC.

Dois modos:

  1) Modelo-base .docx + tabela de placeholders
     python gerar_docx.py --modelo <base.docx> --dados <dados.json> --saida <out.docx>

     dados.json e um dict {placeholder: valor}, ex:
     {"{{NOME_AUTOR}}": "FULANO", "{{NUMERO_CNJ}}": "0600700-28.2025..."}

  2) Markdown estruturado (fallback quando nao ha .docx base pronto)
     python gerar_docx.py --md <plano.md> --saida <out.docx>

Saida padrao: Cambria 12pt em todos os runs. Listas com "a), b), c)" nunca com traco.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, Cm
except ImportError:
    print(json.dumps({"erro": "python-docx nao instalado. Rodar: pip install python-docx"}))
    sys.exit(2)


FONT = "Cambria"
SIZE_PT = 12


def force_font(run) -> None:
    run.font.name = FONT
    run.font.size = Pt(SIZE_PT)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), FONT)


def replace_in_run(run, mapa: dict) -> None:
    text = run.text
    mudou = False
    for k, v in mapa.items():
        if k in text:
            text = text.replace(k, str(v))
            mudou = True
    if mudou:
        run.text = text
        force_font(run)


def replace_in_paragraph(p, mapa: dict) -> None:
    # python-docx quebra placeholders entre runs em alguns casos. Consolidamos primeiro.
    full = "".join(r.text for r in p.runs)
    precisa_merge = any(k in full and not any(k in r.text for r in p.runs) for k in mapa)
    if precisa_merge and p.runs:
        p.runs[0].text = full
        for r in p.runs[1:]:
            r.text = ""
    for r in p.runs:
        replace_in_run(r, mapa)


def process_doc(doc, mapa: dict) -> int:
    count = 0
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapa)
        count += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, mapa)
                    count += 1
    return count


def substituir_tracos_por_letras(doc) -> int:
    """Troca marcadores de lista com traco por a)/b)/c)."""
    conv = 0
    rx = re.compile(r"^\s*[-\u2014\u2013\u2022]\s+")
    for p in doc.paragraphs:
        if p.style.name.lower().startswith("list") or rx.match(p.text or ""):
            if p.runs:
                first = p.runs[0]
                first.text = rx.sub("", first.text, count=1)
                conv += 1
    return conv


def checar_placeholders_residuais(doc) -> list[str]:
    residuais = []
    rx = re.compile(r"\{\{[^}]+\}\}")
    for p in doc.paragraphs:
        for m in rx.finditer(p.text or ""):
            residuais.append(m.group(0))
    return residuais


def gerar_de_modelo(modelo: Path, dados: dict, saida: Path) -> dict:
    doc = Document(modelo)
    total = process_doc(doc, dados)
    substituidos_tracos = substituir_tracos_por_letras(doc)
    residuais = checar_placeholders_residuais(doc)
    doc.save(saida)
    return {
        "modo": "modelo",
        "modelo": str(modelo),
        "saida": str(saida),
        "paragrafos_processados": total,
        "substituicoes_realizadas": len(dados),
        "tracos_convertidos_em_letras": substituidos_tracos,
        "placeholders_residuais": residuais,
    }


def gerar_de_markdown(md: Path, saida: Path) -> dict:
    texto = md.read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(SIZE_PT)
    for linha in texto.splitlines():
        if linha.startswith("# "):
            p = doc.add_heading(linha[2:].strip(), level=1)
        elif linha.startswith("## "):
            p = doc.add_heading(linha[3:].strip(), level=2)
        else:
            p = doc.add_paragraph(linha.rstrip())
        for r in p.runs:
            force_font(r)
    doc.save(saida)
    return {
        "modo": "markdown",
        "md": str(md),
        "saida": str(saida),
    }


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--modelo", type=Path)
    ap.add_argument("--dados", type=Path)
    ap.add_argument("--md", type=Path)
    ap.add_argument("--saida", type=Path, required=True)
    args = ap.parse_args()

    if args.modelo and args.dados:
        dados = json.loads(args.dados.read_text(encoding="utf-8"))
        result = gerar_de_modelo(args.modelo, dados, args.saida)
    elif args.md:
        result = gerar_de_markdown(args.md, args.saida)
    else:
        print(json.dumps({"erro": "use --modelo+--dados ou --md"}))
        return 2
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
