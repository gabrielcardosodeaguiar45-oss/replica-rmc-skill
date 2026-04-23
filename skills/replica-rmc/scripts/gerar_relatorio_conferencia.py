#!/usr/bin/env python3
"""gerar_relatorio_conferencia.py — gera o .docx paralelo do revisor.

Usado pelo subagent revisor-replica-rmc para materializar o relatorio de
conferencia com estrutura padrao (identificacao, tempestividade, conteudo,
erros comuns, layout, fechamento, arquivo, 29 erros herdados, regras aplicadas,
classificacao final).

Uso:
  python gerar_relatorio_conferencia.py --relatorio <json> --saida <out.docx>

O JSON segue o template em references/schema_relatorio.json (a ser criado).
"""
from __future__ import annotations

import argparse
import json
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, Cm
except ImportError:
    print(json.dumps({"erro": "python-docx nao instalado. Rodar: pip install python-docx"}))
    sys.exit(2)


FONT = "Cambria"
SIZE_PT = 11


def force_font(run) -> None:
    run.font.name = FONT
    run.font.size = Pt(SIZE_PT)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), FONT)


def add_p(doc, text: str, bold: bool = False, size: int | None = None) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    force_font(r)
    if size:
        r.font.size = Pt(size)


def add_tabela(doc, titulo: str, itens: list[dict]) -> None:
    add_p(doc, titulo, bold=True, size=13)
    if not itens:
        add_p(doc, "(nenhum item avaliado)")
        return
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Status"
    hdr[2].text = "Observacao"
    for row in hdr:
        for p in row.paragraphs:
            for r in p.runs:
                r.bold = True
                force_font(r)
    for item in itens:
        r = tbl.add_row().cells
        r[0].text = item.get("item", "")
        r[1].text = item.get("status", "")
        r[2].text = item.get("observacao", "") or ""
        for c in r:
            for p in c.paragraphs:
                for rn in p.runs:
                    force_font(rn)


def gerar(relatorio: dict, saida: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)

    add_p(doc, "RELATORIO DE CONFERENCIA — Replica RMC/RCC", bold=True, size=16)
    add_p(doc, "")
    meta = relatorio.get("meta") or {}
    for rotulo, chave in (("Processo", "cnj"), ("Autor", "autor"), ("Banco", "banco"), ("Data", "data")):
        add_p(doc, f"{rotulo}: {meta.get(chave, '') or ''}", bold=False)
    add_p(doc, "")

    blocos = relatorio.get("blocos") or []
    for i, bloco in enumerate(blocos, start=1):
        titulo = f"{i}. {bloco.get('titulo', 'BLOCO')}"
        add_tabela(doc, titulo, bloco.get("itens") or [])
        add_p(doc, "")

    erros = relatorio.get("erros_herdados") or []
    if erros:
        add_tabela(doc, f"{len(blocos) + 1}. ERROS HERDADOS (29 itens)", erros)
        add_p(doc, "")

    regras = relatorio.get("regras") or {}
    if regras:
        add_p(doc, f"{len(blocos) + 2}. REGRAS APLICADAS (das 16)", bold=True, size=13)
        add_p(doc, f"Regras listadas no plano: {regras.get('listadas', [])}")
        add_p(doc, f"Regras refletidas no texto: {regras.get('refletidas', [])}")
        add_p(doc, f"Regras AUSENTES: {regras.get('ausentes', [])}")
        add_p(doc, "")

    classificacao = relatorio.get("classificacao") or {}
    add_p(doc, f"{len(blocos) + 3}. CLASSIFICACAO FINAL", bold=True, size=13)
    add_p(doc, f"Resultado: {classificacao.get('resultado', 'INDETERMINADO')}", bold=True)
    ajustes = classificacao.get("ajustes") or []
    if ajustes:
        add_p(doc, "Ajustes pendentes:", bold=True)
        for idx, aj in enumerate(ajustes, start=1):
            add_p(doc, f"  {idx}. {aj}")
    else:
        add_p(doc, "Sem ajustes pendentes.")

    doc.save(saida)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--relatorio", type=Path, required=True)
    ap.add_argument("--saida", type=Path, required=True)
    args = ap.parse_args()
    relatorio = json.loads(args.relatorio.read_text(encoding="utf-8"))
    gerar(relatorio, args.saida)
    print(json.dumps({"saida": str(args.saida), "data": date.today().isoformat()}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
