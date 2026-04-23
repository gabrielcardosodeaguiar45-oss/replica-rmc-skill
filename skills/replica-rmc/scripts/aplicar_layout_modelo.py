#!/usr/bin/env python3
"""aplicar_layout_modelo.py — mescla o corpo de uma replica dentro do esqueleto
do .docx modelo do vault, preservando timbre (header/footer com imagens),
margens, estilos Heading 1/Heading 2 e aplicando layout padrao do escritorio.

Problema que resolve:
  O `gerar_docx.py` em modo direto cria um .docx do zero — perde timbre do
  escritorio, margens corretas, estilos Heading nativos. Este script recebe:
    - .docx MODELO: o .docx original do vault (ex: "1. Replica PAN.docx")
    - .docx CONTEUDO: a replica redigida (qualquer layout)
  E produz o .docx final com:
    1) Header/footer do modelo (timbre preservado).
    2) Margens do modelo.
    3) Estilos Heading 1 (titulos principais) / Heading 2 (subtitulos DENTRO de
       'DOS FUNDAMENTOS JURIDICOS DOS PEDIDOS') aplicados aos titulos.
    4) Cambria 12pt no corpo; 10pt para citacoes legais/jurisprudenciais.
    5) first_line_indent ~1.5cm e line_spacing 1.5 nos paragrafos Normal.
    6) left_indent ~4cm nos paragrafos de citacao.
    7) Enderecamento, processo no, fecho com alinhamentos corretos.
    8) Uma linha em branco antes de cada Heading 1.

Uso:
  python aplicar_layout_modelo.py --modelo <modelo.docx> \\
         --conteudo <replica_redigida.docx> --saida <final.docx>

Saida:
  JSON com metricas (paragrafos normais, citacoes, headings, bold_removido, ...).
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from copy import deepcopy
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print(json.dumps({"erro": "python-docx nao instalado. pip install python-docx"}))
    sys.exit(2)


FONT = "Cambria"
SIZE_CORPO = Pt(12)
SIZE_CITE = Pt(10)
FIRST_LINE_EMU = 540385    # ~ 1.50 cm (first_line_indent padrao)
JURIS_LEFT_EMU = 1439545   # ~ 4.00 cm (left_indent de citacoes)
LINE_SPACING = 1.5


PATTERNS_CITE = [
    r"^\s*Art\.\s*\d+",
    r"^\s*Par[áa]grafo\s+(?:[úu]nico|primeiro|\d+)",
    r"^\s*§\s*\d+",
    r"^\s*[IVX]+\s*[-–]",
    r"^\s*\d+º?\s*[-–]\s",
    r"^[\"\u201C]",
    r"^\s*APELA[ÇC][ÃA]O\s+C[ÍI]VEL",
    r"^\s*AGRAVO\s+(?:DE\s+INSTRUMENTO|INTERNO|EM\s+RECURSO)",
    r"^\s*RECURSO\s+ESPECIAL",
    r"^\s*HABEAS\s+CORPUS",
    r"^\s*CONCLUS[ÃA]O\s+\d+",
    r"^\s*CL[ÁA]USULA\s+\d+",
    r"^\s*S[ÚU]MULA\s+\d+",
    r"^\s*REsp\s+\d",
    r"^\s*EAREsp\s+\d",
    r"^\s*\(TJ[A-Z]{2}",
    r"^\s*TJ[A-Z]{2}\s*,?\s+(Apela[çc][ãa]o|Agravo)",
    r"^\s*Decidiu\s+o\s+Superior",
    r"^\s*\([A-Z][A-Z\s\-\d\.]+,\s+(?:Rel|Ministra|Ministro)",
]
RX_CITE = re.compile("|".join(PATTERNS_CITE))


def eh_ementa_maiuscula(txt: str) -> bool:
    """Detecta se o paragrafo e uma ementa — MAIUSCULAS predominantes e longo."""
    if len(txt) < 80:
        return False
    alfa = sum(1 for c in txt if c.isalpha())
    if alfa == 0:
        return False
    maius = sum(1 for c in txt if c.isupper())
    return maius / alfa > 0.85


def force_font_run(run_elem, size_pt: Pt = SIZE_CORPO) -> None:
    rPr = run_elem.find(qn("w:rPr"))
    if rPr is None:
        rPr = run_elem.makeelement(qn("w:rPr"), {})
        run_elem.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), FONT)
    val = str(int(size_pt.pt * 2))
    sz = rPr.find(qn("w:sz"))
    if sz is None:
        rPr.append(rPr.makeelement(qn("w:sz"), {qn("w:val"): val}))
    else:
        sz.set(qn("w:val"), val)
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        rPr.append(rPr.makeelement(qn("w:szCs"), {qn("w:val"): val}))
    else:
        szCs.set(qn("w:val"), val)


def remover_bold(run_elem) -> bool:
    rPr = run_elem.find(qn("w:rPr"))
    if rPr is None:
        return False
    alterado = False
    for tag in ("w:b", "w:bCs"):
        el = rPr.find(qn(tag))
        if el is not None:
            rPr.remove(el)
            alterado = True
    return alterado


def forcar_bold(run_elem) -> None:
    rPr = run_elem.find(qn("w:rPr"))
    if rPr is None:
        rPr = run_elem.makeelement(qn("w:rPr"), {})
        run_elem.insert(0, rPr)
    if rPr.find(qn("w:b")) is None:
        rPr.append(rPr.makeelement(qn("w:b"), {}))


def copiar_corpo_para_modelo(modelo: Document, conteudo: Document) -> None:
    """Limpa o body do modelo (exceto sectPrs) e insere os elementos do conteudo."""
    body = modelo.element.body
    sectPr_final = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sectPr_final = child

    removiveis = []
    for child in list(body):
        if child.tag == qn("w:p"):
            pPr = child.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                continue
            removiveis.append(child)
        elif child.tag == qn("w:tbl"):
            removiveis.append(child)
    for r in removiveis:
        body.remove(r)

    conteudo_elems = []
    for e in conteudo.element.body:
        if e.tag == qn("w:sectPr"):
            continue
        if e.tag == qn("w:p"):
            pPr = e.find(qn("w:pPr"))
            if pPr is not None:
                emb = pPr.find(qn("w:sectPr"))
                if emb is not None:
                    pPr.remove(emb)
        conteudo_elems.append(deepcopy(e))

    primeiro_fixo = None
    for child in body:
        if child.tag == qn("w:p"):
            pPr = child.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                primeiro_fixo = child
                break
        if child.tag == qn("w:sectPr"):
            primeiro_fixo = child
            break

    if primeiro_fixo is None:
        for el in conteudo_elems:
            body.append(el)
    else:
        for el in conteudo_elems:
            primeiro_fixo.addprevious(el)


def aplicar_estilos_e_layout(doc: Document) -> dict:
    estilos = {s.name for s in doc.styles}
    if "Heading 1" not in estilos or "Heading 2" not in estilos:
        return {"erro": "estilos Heading 1/Heading 2 ausentes no modelo"}
    h1 = doc.styles["Heading 1"]
    h2 = doc.styles["Heading 2"]

    paragrafos = doc.paragraphs

    idx_fundamentos = None
    idx_pedidos = None
    for i, p in enumerate(paragrafos):
        t = p.text.strip().upper()
        if "FUNDAMENTOS JUR" in t and "PEDIDOS" in t and idx_fundamentos is None:
            idx_fundamentos = i
        elif t == "DOS PEDIDOS" and idx_pedidos is None:
            idx_pedidos = i

    metric = {
        "headings_h1": 0,
        "headings_h2": 0,
        "normais": 0,
        "citacoes": 0,
        "bold_removidos": 0,
        "linhas_em_branco_adicionadas": 0,
        "primeiros_sem_indent": 0,
    }

    contador_prim = 0
    for i, p in enumerate(paragrafos):
        txt = p.text.strip()
        if not txt:
            continue

        # Linha de assinatura/fecho — tratar separadamente
        fecho_data = re.match(r"^[^\d]+,\s+\d{1,2}\s+de\s+\w+\s+de\s+\d{4}\.?$", txt)
        eh_fecho = (
            txt.lower().startswith("nestes termos")
            or fecho_data is not None
            or txt.startswith("___")
            or ("OAB" in txt and "/" in txt)
        )
        # Nome do advogado (palavra CAPITALIZADA curta, nao CAPS, sem contexto processual)
        eh_nome_adv = (
            not txt.isupper()
            and len(txt.split()) <= 6
            and sum(1 for w in txt.split() if w and w[0].isupper()) >= 2
            and not any(kw in txt.lower() for kw in ["requerida", "autor", "excel", "processo"])
            and i >= len(paragrafos) - 10  # so nas ultimas 10 linhas
        )

        # 1) Cabecalho (primeiro paragrafo com "EXCELENT")
        if i == 0 and txt.startswith("EXCELENT"):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.first_line_indent = None
            pf.left_indent = None
            pf.line_spacing = LINE_SPACING
            for r in p.runs:
                force_font_run(r._element, SIZE_CORPO)
                forcar_bold(r._element)
            continue

        # 2) "Processo no ..."
        if txt.lower().startswith("processo"):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pf = p.paragraph_format
            pf.first_line_indent = None
            pf.left_indent = None
            pf.line_spacing = LINE_SPACING
            for r in p.runs:
                force_font_run(r._element, SIZE_CORPO)
                forcar_bold(r._element)
            continue

        # 3) Fecho: centralizar + sem indent
        if eh_fecho:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.first_line_indent = None
            pf.left_indent = None
            pf.line_spacing = LINE_SPACING
            for r in p.runs:
                force_font_run(r._element, SIZE_CORPO)
                if "OAB" in txt:
                    forcar_bold(r._element)
            continue

        if eh_nome_adv:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.first_line_indent = None
            pf.left_indent = None
            pf.line_spacing = LINE_SPACING
            for r in p.runs:
                force_font_run(r._element, SIZE_CORPO)
                forcar_bold(r._element)
            continue

        # 4) Titulo: CAPS + texto >= 6 alfa + (center OU bold no run original)
        alfa = sum(1 for c in txt if c.isalpha())
        is_upper = txt == txt.upper() and alfa >= 6
        any_bold = any(r.bold for r in p.runs if r.text.strip())

        if is_upper and (any_bold or p.alignment == WD_ALIGN_PARAGRAPH.CENTER) and not eh_ementa_maiuscula(txt):
            # Limiar: ementa tem > 85% caps E > 80 chars — titulos reais tem ate ~100 chars
            # Distinguir: se comeca com CONCLUSAO, CLAUSULA, SUMULA → citacao
            if RX_CITE.match(txt):
                # e citacao
                pass
            else:
                if idx_fundamentos is not None and idx_pedidos is not None and idx_fundamentos < i < idx_pedidos:
                    p.style = h2
                    metric["headings_h2"] += 1
                else:
                    p.style = h1
                    metric["headings_h1"] += 1
                for r in p.runs:
                    if remover_bold(r._element):
                        metric["bold_removidos"] += 1
                    force_font_run(r._element, SIZE_CORPO)
                continue

        # 5) Detectar citacao
        tem_fonte_menor = any(
            r.font.size and r.font.size.pt <= 10.5 for r in p.runs if r.text.strip()
        )
        is_cite = (
            bool(RX_CITE.match(txt))
            or tem_fonte_menor
            or (eh_ementa_maiuscula(txt) and i > 3)
        )

        pf = p.paragraph_format
        if is_cite:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.first_line_indent = None
            pf.left_indent = Emu(JURIS_LEFT_EMU)
            pf.line_spacing = 1.0
            for r in p.runs:
                force_font_run(r._element, SIZE_CITE)
                if remover_bold(r._element):
                    metric["bold_removidos"] += 1
            metric["citacoes"] += 1
            continue

        # 6) Paragrafo normal de corpo
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = LINE_SPACING

        contador_prim += 1
        if contador_prim <= 1:  # so qualificacao do autor (apos enderecamento e processo que ja foram tratados)
            pf.first_line_indent = None
            pf.left_indent = None
            metric["primeiros_sem_indent"] += 1
        else:
            pf.first_line_indent = Emu(FIRST_LINE_EMU)
            pf.left_indent = None
        for r in p.runs:
            force_font_run(r._element, SIZE_CORPO)
            if remover_bold(r._element):
                metric["bold_removidos"] += 1
        metric["normais"] += 1

    # 7) Garantir linha em branco antes de cada Heading 1
    paragrafos = doc.paragraphs
    to_insert = []
    for i, p in enumerate(paragrafos):
        if p.style.name == "Heading 1" and i > 0:
            prev = paragrafos[i - 1]
            if prev.text.strip():
                to_insert.append(i)

    for idx in reversed(to_insert):
        h = doc.paragraphs[idx]
        novo = deepcopy(doc.paragraphs[0]._element)
        for c in list(novo):
            novo.remove(c)
        h._element.addprevious(novo)
        metric["linhas_em_branco_adicionadas"] += 1

    # 8) Numeracao APENAS em Heading 1 (algarismos romanos).
    # Heading 2 NAO recebe numeracao manual (o Word pode ter auto-numbering
    # pelo estilo, o que causa "g. g)" duplicado). Alem disso, removemos
    # qualquer numPr existente para desligar auto-numbering completamente.
    metric["h1_numerados"] = 0
    metric["h2_limpos"] = 0
    metric["numPr_removidos"] = 0
    cont_h1 = 0
    rx_roman = re.compile(r"^[IVXLCDM]+\s*[—\-–]\s*")
    rx_letra = re.compile(r"^[a-z]\)\s*", re.IGNORECASE)

    # Primeiro: desligar auto-numbering dos estilos Heading 1 e Heading 2
    for style_name in ("Heading 1", "Heading 2"):
        try:
            style = doc.styles[style_name]
            pPr_style = style.element.find(qn("w:pPr"))
            if pPr_style is not None:
                numPr = pPr_style.find(qn("w:numPr"))
                if numPr is not None:
                    pPr_style.remove(numPr)
                    metric["numPr_removidos"] += 1
        except KeyError:
            pass

    for p in doc.paragraphs:
        # Remover qualquer numPr embutido em paragrafos de heading
        if p.style.name in ("Heading 1", "Heading 2"):
            pPr = p._element.find(qn("w:pPr"))
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                if numPr is not None:
                    pPr.remove(numPr)
                    metric["numPr_removidos"] += 1

        if p.style.name == "Heading 1":
            cont_h1 += 1
            t = p.text.strip()
            t = rx_roman.sub("", t)
            t = corrigir_acentos_titulo(t)  # trava anti-sem-acento
            novo = f"{_int_to_roman(cont_h1)} — {t}"
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            p.add_run(novo)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                force_font_run(r._element, SIZE_CORPO)
            metric["h1_numerados"] += 1
        elif p.style.name == "Heading 2":
            # Garantir que nao ha numeracao manual "a) ", "b) " etc +
            # restaurar acentos
            t = p.text.strip()
            t_limpo = rx_letra.sub("", t)
            t_acentuado = corrigir_acentos_titulo(t_limpo)
            if t != t_acentuado:
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
                p.add_run(t_acentuado)
                for r in p.runs:
                    force_font_run(r._element, SIZE_CORPO)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                metric["h2_limpos"] += 1

    return metric


def _int_to_roman(n: int) -> str:
    tabela = [
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    ]
    out = ""
    for val, sym in tabela:
        while n >= val:
            out += sym
            n -= val
    return out


# Dicionario de correcao de acentos para titulos em CAIXA ALTA.
# Causa recorrente: modelos antigos do escritorio tem titulos sem acento; ao copiar o
# redator herda. Aqui restauramos.
ACENTOS_TITULOS = {
    "SINTESE": "SÍNTESE", "CONTESTACAO": "CONTESTAÇÃO", "INEPCIA": "INÉPCIA",
    "AUSENCIA": "AUSÊNCIA", "MINIMA": "MÍNIMA", "DELIMITACAO": "DELIMITAÇÃO",
    "CONTROVERSIA": "CONTROVÉRSIA", "CARENCIA": "CARÊNCIA", "ACAO": "AÇÃO",
    "CONFIRMACAO": "CONFIRMAÇÃO", "PROCURACAO": "PROCURAÇÃO",
    "LITIGANCIA": "LITIGÂNCIA", "PREDATORIA": "PREDATÓRIA",
    "RECOMENDACAO": "RECOMENDAÇÃO", "JUSTICA": "JUSTIÇA",
    "PRESCRICAO": "PRESCRIÇÃO", "DECADENCIA": "DECADÊNCIA",
    "JURIDICOS": "JURÍDICOS", "JURIDICA": "JURÍDICA",
    "INSTRUCAO": "INSTRUÇÃO", "CONTRATACAO": "CONTRATAÇÃO",
    "CONTRATACOES": "CONTRATAÇÕES", "DOSSIE": "DOSSIÊ",
    "TECNICO": "TÉCNICO", "ESPECIFICO": "ESPECÍFICO",
    "FORMALIZACAO": "FORMALIZAÇÃO", "BANCARIO": "BANCÁRIO",
    "RESIDENCIA": "RESIDÊNCIA", "DIVERGENCIAS": "DIVERGÊNCIAS",
    "PARAMETROS": "PARÂMETROS", "CARTAO": "CARTÃO",
    "EMPRESTIMO": "EMPRÉSTIMO", "CREDITO": "CRÉDITO",
    "VICIO": "VÍCIO", "INFORMACAO": "INFORMAÇÃO",
    "PRATICA": "PRÁTICA", "SUMULA": "SÚMULA",
    "ILICITO": "ILÍCITO", "RESTITUICAO": "RESTITUIÇÃO",
    "INVERSAO": "INVERSÃO", "ONUS": "ÔNUS",
    "HONORARIOS": "HONORÁRIOS", "ADVOCATICIOS": "ADVOCATÍCIOS",
    "BENEFICIO": "BENEFÍCIO", "INERCIA": "INÉRCIA",
    "RENUNCIA": "RENÚNCIA", "INDEXACAO": "INDEXAÇÃO",
    "CONFISSAO": "CONFISSÃO", "PROPRIO": "PRÓPRIO",
    "ALEGACAO": "ALEGAÇÃO", "DIVIDA": "DÍVIDA",
    "MINIMO": "MÍNIMO", "TRAMITACAO": "TRAMITAÇÃO",
    "ANUENCIA": "ANUÊNCIA", "COMPENSACAO": "COMPENSAÇÃO",
    "IMPUGNACAO": "IMPUGNAÇÃO", "Nº": "nº",
    # crases que aparecem em titulos
    "QUANTO A ALEGACAO": "QUANTO À ALEGAÇÃO",
    "QUANTO A MODALIDADE": "QUANTO À MODALIDADE",
    "QUANTO A CONTRATACAO": "QUANTO À CONTRATAÇÃO",
}


def corrigir_acentos_titulo(texto: str) -> str:
    """Aplica dicionario de acentos no texto de um titulo. Idempotente."""
    novo = texto
    for k in sorted(ACENTOS_TITULOS.keys(), key=len, reverse=True):
        padrao = re.escape(k) if " " in k else r"\b" + re.escape(k) + r"\b"
        if re.search(padrao, novo):
            novo = re.sub(padrao, ACENTOS_TITULOS[k], novo)
    return novo


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--modelo", type=Path, required=True, help=".docx modelo (com timbre) do vault")
    ap.add_argument("--conteudo", type=Path, required=True, help=".docx com texto redigido pelo redator")
    ap.add_argument("--saida", type=Path, required=True, help=".docx final")
    args = ap.parse_args()

    if not args.modelo.exists():
        print(json.dumps({"erro": f"modelo nao encontrado: {args.modelo}"}))
        return 2
    if not args.conteudo.exists():
        print(json.dumps({"erro": f"conteudo nao encontrado: {args.conteudo}"}))
        return 2

    # Copiar modelo para saida (preserva timbre, relationships)
    shutil.copy2(args.modelo, args.saida)

    destino = Document(args.saida)
    conteudo = Document(args.conteudo)

    copiar_corpo_para_modelo(destino, conteudo)
    destino.save(args.saida)

    # Reabrir e aplicar layout
    destino = Document(args.saida)
    metric = aplicar_estilos_e_layout(destino)
    destino.save(args.saida)

    # Verificar resultado
    verif = Document(args.saida)
    imgs_header = 0
    imgs_footer = 0
    for s in verif.sections:
        for pp in s.header.paragraphs:
            imgs_header += len(pp._element.findall(".//" + qn("w:drawing")))
        for pp in s.footer.paragraphs:
            imgs_footer += len(pp._element.findall(".//" + qn("w:drawing")))

    result = {
        "saida": str(args.saida),
        "timbre_header_imgs": imgs_header,
        "timbre_footer_imgs": imgs_footer,
        "secoes": len(verif.sections),
        "paragrafos_corpo": len(verif.paragraphs),
        "metric": metric,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
