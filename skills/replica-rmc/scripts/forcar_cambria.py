#!/usr/bin/env python3
"""forcar_cambria.py — forca fonte Cambria em todo o .docx.

Erro 19 recorrente: colagem de textos do vault ou do modelo traz Segoe UI, Arial
ou Calibri em runs especificos. Este script percorre todos os paragrafos e runs,
zera font.name e aplica "Cambria" + rFonts para ASCII/hAnsi/eastAsia.

Uso:
  python forcar_cambria.py <caminho.docx>

Saida:
  {"arquivo": "...", "runs_afetados": N, "tamanho_padrao_pt": 12}
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError:
    print(json.dumps({"erro": "python-docx nao instalado. Rodar: pip install python-docx"}))
    sys.exit(2)


FONT = "Cambria"
SIZE_PT = 12


def force_font_run(run, font_name: str = FONT, size_pt: int = SIZE_PT) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)


def walk(doc) -> int:
    count = 0
    for p in doc.paragraphs:
        for r in p.runs:
            force_font_run(r)
            count += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        force_font_run(r)
                        count += 1
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header is None:
                continue
            for p in header.paragraphs:
                for r in p.runs:
                    force_font_run(r)
                    count += 1
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer is None:
                continue
            for p in footer.paragraphs:
                for r in p.runs:
                    force_font_run(r)
                    count += 1
    return count


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("docx", type=Path)
    args = ap.parse_args()
    if not args.docx.exists():
        print(json.dumps({"erro": f"arquivo nao encontrado: {args.docx}"}))
        return 2
    doc = Document(args.docx)
    afetados = walk(doc)
    doc.save(args.docx)
    print(json.dumps(
        {
            "arquivo": str(args.docx),
            "runs_afetados": afetados,
            "tamanho_padrao_pt": SIZE_PT,
            "fonte": FONT,
        },
        ensure_ascii=False,
        indent=2,
    ))
    return 0


if __name__ == "__main__":
    sys.exit(main())
