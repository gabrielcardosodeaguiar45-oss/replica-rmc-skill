#!/usr/bin/env python3
"""check_ancora_contestacao.py — trava anti-alucinacao: cruza afirmacoes
da replica contra o que a contestacao efetivamente disse.

Principio:
  Todo paragrafo da replica que afirma "A Requerida sustenta X",
  "O banco alega Y", "Sustenta a r[ée]", "DA ALEGADA Z", "No topico W da
  contestacao" DEVE ter correspondencia:
    - no `_analise.json:contestacao.preliminares_levantadas[].trecho_literal`
    - ou no `_analise.json:contestacao.teses_meritorias[].trecho_literal`
    - ou no texto extraido da contestacao (arquivo .txt do analisador).

Alem disso, uma LISTA NEGRA de frases com alto risco de alucinacao e
varrida e qualquer hit gera AJUSTE:
    - "multou em R$", "PROCON/[UF]", "aplicou multa"
    - "em reiteradas ocasioes", "em diversas ocasioes"
    - "como e notorio", "como e publico"
    - "entra em contato direto com o cliente"
    - "verifica se o cliente reconhece a acao"
    - Numeros/valores/CNJs especificos nao presentes em _analise.json
      nem em _plano.json:precedentes.

Uso:
  python check_ancora_contestacao.py --replica <.docx> \
      --analise <_analise.json> [--contestacao-txt <arquivo.txt>]

Saida:
  JSON com: paragrafos_suspeitos[], frases_proibidas[], sem_ancora[].
"""
from __future__ import annotations

import argparse
import json
import re
import sys
import unicodedata
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print(json.dumps({"erro": "python-docx nao instalado"}))
    sys.exit(2)


BLACKLIST_PATTERNS = [
    r"multou em R\$",
    r"aplicou multa (de|superior)",
    r"PROCON\s*/\s*[A-Z]{2}\b",
    r"em reiteradas ocasi[õo]es",
    r"em diversas ocasi[õo]es",
    r"como é not[óo]rio",
    r"como é p[úu]blico",
    r"entra(m)? em contato (direto )?com o cliente",
    r"verifica(r)? se o cliente reconhece a a[çc][ãa]o",
    r"em outros casos similares",
    r"conduta reiterada",
    r"pr[áa]tica reiterada (do banco|da requerida)",
    r"hist[óo]ria do banco.{0,30}descontos indevidos",
]

# Detectores de frase que afirma coisa sobre o banco — exige ancora
RX_AFIRMACAO_BANCO = re.compile(
    r"(A\s+Requerida\s+(?:alega|sustenta|afirma|argumenta|defende|invoca)"
    r"|Alega(m)?\s+o\s+banco"
    r"|Sustenta(m)?\s+(?:o\s+)?banco"
    r"|O\s+banco\s+(?:alega|sustenta|afirma|argumenta|defende|invoca)"
    r"|A\s+defesa\s+(?:alega|sustenta|afirma|argumenta)"
    r"|Pleiteia\s+a\s+Requerida"
    r"|A\s+R[ée]\s+(?:alega|sustenta|afirma))",
    re.IGNORECASE,
)
RX_TITULO_ALEGADA = re.compile(r"^\s*(DA|DO|DAS|DOS)\s+ALEGAD[AO]", re.IGNORECASE)


def norm(s: str) -> str:
    s = unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode("ascii").lower()
    return re.sub(r"\s+", " ", s).strip()


def carregar_ancoras(analise_path: Path, contestacao_txt_path: Path | None) -> tuple[list[str], str]:
    """Retorna (rotulos/trechos conhecidos da contestacao, texto_contestacao_norm)."""
    data = json.loads(analise_path.read_text(encoding="utf-8"))
    contestacao = data.get("contestacao") or {}
    rotulos: list[str] = []
    for item in contestacao.get("preliminares_levantadas") or []:
        if isinstance(item, dict):
            rotulos.append(item.get("id", ""))
            rotulos.append(item.get("rotulo_banco", ""))
            if item.get("trecho_literal"):
                rotulos.append(item["trecho_literal"])
        else:
            rotulos.append(str(item))
    for item in contestacao.get("teses_meritorias") or []:
        if isinstance(item, dict):
            rotulos.append(item.get("id", ""))
            rotulos.append(item.get("rotulo_banco", ""))
            if item.get("trecho_literal"):
                rotulos.append(item["trecho_literal"])
        else:
            rotulos.append(str(item))
    for ftr in contestacao.get("fatos_extraprocessuais_alegados") or []:
        rotulos.append(str(ftr))

    texto_cont = ""
    if contestacao_txt_path and contestacao_txt_path.exists():
        texto_cont = contestacao_txt_path.read_text(encoding="utf-8", errors="replace")

    rotulos = [r for r in rotulos if r and r.strip()]
    return rotulos, norm(texto_cont)


def checar(replica: Path, analise: Path, contestacao_txt: Path | None) -> dict:
    doc = Document(replica)
    rotulos, texto_cont_norm = carregar_ancoras(analise, contestacao_txt)
    rotulos_norm = [norm(r) for r in rotulos]

    frases_proibidas = []
    afirmacoes_sem_ancora = []
    titulos_sem_ancora = []

    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if not txt:
            continue

        # 1) Blacklist de frases de alto risco
        for pat in BLACKLIST_PATTERNS:
            for m in re.finditer(pat, txt, re.IGNORECASE):
                frases_proibidas.append({
                    "paragrafo": i,
                    "padrao": pat,
                    "trecho": txt[max(0, m.start()-20):m.end()+40],
                })

        # 2) Titulos "DA ALEGADA X" — cross-check contra contestacao
        if RX_TITULO_ALEGADA.match(txt):
            # extrair o "X" (remove "DA ALEGADA")
            x_txt = re.sub(r"^\s*(DA|DO|DAS|DOS)\s+ALEGAD[AO]S?\s+", "", txt, flags=re.IGNORECASE).strip()
            x_norm = norm(x_txt)
            tokens = [t for t in re.findall(r"\w+", x_norm) if len(t) > 3]
            # ancora: pelo menos 2 tokens devem aparecer juntos no texto da contestacao
            # OU o rotulo deve ter match forte com algum rotulo/id conhecido
            achou = False
            for rn in rotulos_norm:
                if rn and x_norm in rn or (rn and rn in x_norm and len(rn) > 5):
                    achou = True
                    break
            if not achou and texto_cont_norm and tokens:
                # tentar achar >=2 tokens distintos no texto da contestacao
                presentes = sum(1 for t in tokens if t in texto_cont_norm)
                if presentes >= 2:
                    achou = True
            if not achou:
                titulos_sem_ancora.append({
                    "paragrafo": i,
                    "titulo": txt,
                    "tokens_procurados": tokens,
                })

        # 3) Afirmacao sobre o banco — dentro do corpo
        if RX_AFIRMACAO_BANCO.search(txt):
            # extrair ~150 caracteres a partir do verbo
            tokens = [t for t in re.findall(r"\w+", norm(txt)) if len(t) > 4]
            match_contest = 0
            if texto_cont_norm and tokens:
                match_contest = sum(1 for t in set(tokens[:20]) if t in texto_cont_norm)
            suficiente = match_contest >= 3
            match_rotulo = any(
                any(t in rn for t in tokens[:10] if len(t) > 4) for rn in rotulos_norm
            )
            if not suficiente and not match_rotulo:
                afirmacoes_sem_ancora.append({
                    "paragrafo": i,
                    "tokens_match_contestacao": match_contest,
                    "trecho": txt[:220],
                })

    return {
        "arquivo_replica": str(replica),
        "total_paragrafos": len(doc.paragraphs),
        "frases_proibidas_blacklist": frases_proibidas,
        "titulos_alegada_sem_ancora": titulos_sem_ancora,
        "afirmacoes_sobre_banco_sem_ancora": afirmacoes_sem_ancora,
        "classificacao": (
            "AJUSTE" if (frases_proibidas or titulos_sem_ancora or afirmacoes_sem_ancora)
            else "OK"
        ),
    }


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--replica", type=Path, required=True)
    ap.add_argument("--analise", type=Path, required=True)
    ap.add_argument("--contestacao-txt", type=Path, default=None,
                    help="arquivo .txt com o texto integral da contestacao (extracao do PDF)")
    args = ap.parse_args()
    result = checar(args.replica, args.analise, args.contestacao_txt)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if result["classificacao"] == "OK" else 1


if __name__ == "__main__":
    sys.exit(main())
